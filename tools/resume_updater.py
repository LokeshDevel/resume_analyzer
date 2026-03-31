import os
import io
import traceback
from pathlib import Path
from docx import Document

def update_docx_in_place(file_path: str, analysis_data: dict) -> bytes:
    """
    Reads an original DOCX file, intelligently replaces the user's original
    bullet points with the AI's improved bullet points, and returns
    a byte stream of the modified document (preserving template).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Original file not found at {file_path}")

    doc = Document(file_path)
    
    # Map original strings to improved strings
    replacements = {}
    for b in analysis_data.get("improved_bullet_points", []):
        orig = b.get("original", "").strip()
        improved = b.get("improved", "").strip()
        if orig and improved:
            replacements[orig] = improved

    # Perform replacement
    # We iterate over paragraphs, and if an exact (or highly similar) match
    # is found, we swap the text.
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not p_text:
            continue
            
        for orig, improved in replacements.items():
            # A strict sub-string match is used. 
            # Note: docx `p.text` often loses trailing/leading spaces, so we strip both.
            if len(orig) > 10 and (orig in p_text or orig == p_text):
                # Replace the text.
                new_text = p_text.replace(orig, improved)
                
                # To preserve the paragraph style (e.g., bullet point format, font family),
                # we clear all runs except the first, and assign the new text to the first.
                if len(p.runs) > 0:
                    p.runs[0].text = new_text
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    p.add_run(new_text)
                    
                # Once patched, remove from the replacement dict to avoid double-processing
                # (though harmless if it happens)
                break

    # Also iterate over tables! Many resumes use tables for structural layout.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p_text = p.text.strip()
                    if not p_text:
                        continue
                    
                    for orig, improved in replacements.items():
                        if len(orig) > 10 and (orig in p_text or orig == p_text):
                            new_text = p_text.replace(orig, improved)
                            if len(p.runs) > 0:
                                p.runs[0].text = new_text
                                for run in p.runs[1:]:
                                    run.text = ""
                            else:
                                p.add_run(new_text)
                            break

    # Save to a memory buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.read()
