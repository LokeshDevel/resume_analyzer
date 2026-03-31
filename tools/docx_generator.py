"""
docx_generator.py — Generate a structured DOCX report from resume analysis JSON.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
import io

def generate_docx_report(data: dict) -> bytes:
    doc = Document()
    
    # Title
    title = doc.add_heading('Resume Analyzer Report', 0)
    title.alignment = 1 # center
    
    # Overview
    doc.add_heading('Overview', level=1)
    p = doc.add_paragraph()
    p.add_run(f"ATS Score: {data.get('ats_score', 0)}/100\n").bold = True
    p.add_run(f"Score Band: {data.get('score_band', '')}\n")
    p.add_run(f"Analyzed File: {data.get('original_filename', 'resume')}\n")
    
    # AI Summary
    if data.get('rewritten_summary'):
        doc.add_heading('AI-Rewritten Summary', level=1)
        doc.add_paragraph(data['rewritten_summary'])
    
    # Category Scores
    doc.add_heading('Category Breakdown', level=1)
    for cat_key, cat_val in data.get('category_scores', {}).items():
        cat_name = cat_key.replace('_', ' ').title()
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{cat_name}: {cat_val.get('score', 0)}/100 ").bold = True
        p.add_run(f"(Weight: {cat_val.get('weight', 0)}%) - {cat_val.get('details', '')}")
    
    # Keywords
    doc.add_heading('Keyword Analysis', level=1)
    p = doc.add_paragraph()
    p.add_run("Matched Keywords:\n").bold = True
    matched = data.get('matched_keywords', [])
    p.add_run(", ".join(matched) if matched else "None")
    
    p = doc.add_paragraph()
    p.add_run("Missing Keywords:\n").bold = True
    missing = data.get('missing_keywords', [])
    p.add_run(", ".join(missing) if missing else "None")
    
    # Strengths / Weaknesses
    doc.add_heading('Strengths & Weaknesses', level=1)
    doc.add_paragraph("Strengths:", style='Heading 2')
    for s in data.get('strengths', []):
        doc.add_paragraph(s, style='List Bullet')
        
    doc.add_paragraph("Areas to Improve:", style='Heading 2')
    for w in data.get('weaknesses', []):
        doc.add_paragraph(w, style='List Bullet')
        
    # Recommendations
    recs = data.get('recommendations', [])
    if recs:
        doc.add_heading('Actionable Recommendations', level=1)
        for r in recs:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{r.get('priority', 'medium').upper()}] ").bold = True
            p.add_run(r.get('suggestion', ''))
            
    # Bullet points
    bullets = data.get('improved_bullet_points', [])
    if bullets:
        doc.add_heading('Improved Bullet Points', level=1)
        for b in bullets:
            doc.add_paragraph()
            p1 = doc.add_paragraph()
            p1.add_run("Original: ").bold = True
            p1.add_run(b.get('original', ''))
            
            p2 = doc.add_paragraph()
            p2.add_run("Improved: ").bold = True
            p2.add_run(b.get('improved', ''))
            
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.read()
