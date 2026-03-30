"""
parse_resume.py — Extract raw text from PDF, DOCX, or image files.

Strategy:
  1. PDF  → PyPDF2 direct extraction, fallback to OCR.Space
  2. DOCX → python-docx paragraph extraction
  3. Image → OCR.Space directly

Usage:
    from tools.parse_resume import extract_text
    result = extract_text("uploads/resume.pdf", ocr_api_key="...")
"""

import os
import sys
import json
import requests
from pathlib import Path


def extract_from_pdf(file_path: str) -> dict:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())

        full_text = "\n\n".join(pages)
        return {
            "success": True,
            "text": full_text,
            "method": "direct_pdf",
            "page_count": len(reader.pages),
            "error": None,
        }
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "method": "direct_pdf",
            "page_count": 0,
            "error": str(e),
        }


def extract_from_docx(file_path: str) -> dict:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        full_text = "\n".join(paragraphs)
        return {
            "success": True,
            "text": full_text,
            "method": "direct_docx",
            "page_count": 1,  # DOCX doesn't have page concept in python-docx
            "error": None,
        }
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "method": "direct_docx",
            "page_count": 0,
            "error": str(e),
        }


def extract_via_ocr(file_path: str, api_key: str) -> dict:
    """Extract text from file using OCR.Space API."""
    if not api_key:
        return {
            "success": False,
            "text": "",
            "method": "ocr_space",
            "page_count": 0,
            "error": "OCR_SPACE_API_KEY not provided",
        }

    try:
        with open(file_path, "rb") as f:
            response = requests.post(
                "https://api.ocr.space/parse/image",
                files={"file": (Path(file_path).name, f)},
                data={
                    "apikey": api_key,
                    "language": "eng",
                    "isOverlayRequired": False,
                    "detectOrientation": True,
                    "scale": True,
                    "OCREngine": 2,  # Engine 2 is better for documents
                },
                timeout=60,
            )

        if response.status_code != 200:
            return {
                "success": False,
                "text": "",
                "method": "ocr_space",
                "page_count": 0,
                "error": f"OCR.Space returned HTTP {response.status_code}",
            }

        result = response.json()

        if result.get("IsErroredOnProcessing"):
            error_msg = result.get("ErrorMessage", ["Unknown error"])
            return {
                "success": False,
                "text": "",
                "method": "ocr_space",
                "page_count": 0,
                "error": f"OCR processing error: {error_msg}",
            }

        parsed_results = result.get("ParsedResults", [])
        if not parsed_results:
            return {
                "success": False,
                "text": "",
                "method": "ocr_space",
                "page_count": 0,
                "error": "No parsed results returned from OCR.Space",
            }

        pages = [pr.get("ParsedText", "") for pr in parsed_results]
        full_text = "\n\n".join(pages).strip()

        return {
            "success": True,
            "text": full_text,
            "method": "ocr_space",
            "page_count": len(parsed_results),
            "error": None,
        }

    except requests.Timeout:
        return {
            "success": False,
            "text": "",
            "method": "ocr_space",
            "page_count": 0,
            "error": "OCR.Space request timed out (60s)",
        }
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "method": "ocr_space",
            "page_count": 0,
            "error": str(e),
        }


def extract_text(file_path: str, ocr_api_key: str = None) -> dict:
    """
    Main entry point: extract text from a resume file.

    Strategy:
    - PDF: Try direct extraction first, fall back to OCR if text < 50 chars
    - DOCX: Direct extraction only
    - Image (PNG/JPG): OCR directly

    Returns:
        {
            "success": bool,
            "text": str,
            "method": str,  # "direct_pdf", "direct_docx", "ocr_space"
            "page_count": int,
            "error": str | None
        }
    """
    path = Path(file_path)

    if not path.exists():
        return {
            "success": False,
            "text": "",
            "method": "none",
            "page_count": 0,
            "error": f"File not found: {file_path}",
        }

    ext = path.suffix.lower()

    if ext == ".pdf":
        # Try direct extraction first
        result = extract_from_pdf(file_path)
        if result["success"] and len(result["text"].strip()) >= 50:
            return result
        # Fallback to OCR
        print(f"  ℹ️  Direct PDF extraction yielded {len(result['text'])} chars, trying OCR...")
        return extract_via_ocr(file_path, ocr_api_key)

    elif ext == ".docx":
        return extract_from_docx(file_path)

    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"):
        return extract_via_ocr(file_path, ocr_api_key)

    else:
        # Try reading as plain text
        try:
            text = path.read_text(encoding="utf-8")
            return {
                "success": True,
                "text": text,
                "method": "plain_text",
                "page_count": 1,
                "error": None,
            }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "method": "none",
                "page_count": 0,
                "error": f"Unsupported file format: {ext}. Error: {e}",
            }


def main():
    """CLI usage for testing."""
    if len(sys.argv) < 2:
        print("Usage: python parse_resume.py <file_path>")
        sys.exit(1)

    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent.parent / ".env")

    file_path = sys.argv[1]
    api_key = os.getenv("OCR_SPACE_API_KEY")

    result = extract_text(file_path, api_key)
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
